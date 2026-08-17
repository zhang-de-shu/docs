# -*- coding: utf-8 -*-
"""
draw_utils.py — 专利技术交底书附图绘制工具库（白底 · 纯黑 · 粗线 · 大字）。

引擎核心移植自开源 cnipa-patent-writer（fnjialun，MIT License）的 make_figures.py；
扩展图型（tree/sequence/data_flow/callout/six_views）为同风格自研移植。
依赖：pip install matplotlib（并需一个中文字体）。

设计目标（逐条对应旧版返工病因）：
  ① 纯黑：线 lw≥1.8、色 black、白底保存——细彩线缩放后发灰像"没黑透"。
  ② 不溢出：框高随换行后文字自适应、文字按显示宽度换行到框内（留余量），绝不超出框。
  ③ 竖排单列：流程/模块一律单列自上而下——满版宽框、字大；线只在框间直走，不穿框。
  ④ 不压线：箭头/分组/回流标签一律白底或置于留白；不用竖排单字标签。
  ⑤ 图文匹配：画布宽贴近页宽（6.4in），嵌入（~15cm）后近 1:1，字号≈所设 fs。

用法（高层，推荐；每个函数自带画布并保存 PNG）：
    import draw_utils as du
    du.setup_cjk()   # 或 du.setup_cjk('/path/字体.ttc')
    du.vflow('fig1_flow.png', ["采集", ("处理", "对采集数据做……"), "输出"],
             down_labels={1: "满足条件"}, title="图1 整体流程示意图")
    du.vmodules('fig2_arch.png', [("A模块", "职责……"), ("B模块", "职责……")],
                groups=[("子系统", 0, 1)], feedback=(1, 0, "结果回流"),
                title="图2 系统模块框图")
    du.draw_tree / draw_sequence / draw_data_flow / draw_component_callout / draw_six_views 见后文。

**每画完一张必须 Read 该 PNG，对照 prompts/09-figure-check.md 的自查清单逐条核对后再嵌入。**
"""
import os
import glob
import math
import matplotlib
matplotlib.use("Agg")
import matplotlib.font_manager as fm
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle, FancyArrowPatch

K = "black"
FONT = None

_CANDIDATES = [
    # macOS 系统字体（优先）
    "/System/Library/Fonts/STHeiti Medium.ttc",
    "/System/Library/Fonts/Hiragino Sans GB.ttc",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
    "/System/Library/Fonts/PingFang.ttc",
    # 用户目录
    os.path.expanduser("~/.fonts/NotoSansSC.ttf"),
    os.path.expanduser("~/.fonts/simsun.ttc"),
    os.path.expanduser("~/.fonts/simhei.ttf"),
    # Linux
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
]


def setup_cjk(font_path=None):
    """注册中文字体并设 rcParams，返回字体名。找不到会抛错（中文必豆腐，必须解决）。"""
    global FONT
    path = font_path or os.environ.get("FIG_FONT_PATH")
    if path is None:
        for c in _CANDIDATES:
            if os.path.exists(c):
                path = c
                break
    if path is None:
        for pat in ("/usr/share/fonts/**/*CJK*.ttc", "/System/Library/Fonts/**/*.ttc"):
            hit = glob.glob(pat, recursive=True)
            if hit:
                path = hit[0]
                break
    if path is None or not os.path.exists(path):
        raise RuntimeError("未找到中文字体。请安装 Noto CJK 或显式传 font_path / 设 FIG_FONT_PATH。")
    fm.fontManager.addfont(path)
    FONT = fm.FontProperties(fname=path).get_name()
    plt.rcParams["font.sans-serif"] = [FONT]
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.facecolor"] = "white"
    return FONT


def _ensure_font():
    if FONT is None:
        setup_cjk()


# ---------- 文本换行（按显示宽度，CJK 计 2） ----------
def disp_w(s):
    return sum(2 if ord(c) > 0x2E7F else 1 for c in str(s))


def wrap_cjk(text, width):
    """按显示宽度换行（显式 \\n 也断行），返回行列表。width 为每行显示宽度上限。"""
    out, cur, w = [], "", 0
    for ch in str(text):
        if ch == "\n":
            out.append(cur); cur, w = "", 0; continue
        cw = 2 if ord(ch) > 0x2E7F else 1
        if w + cw > width and cur:
            out.append(cur); cur, w = ch, cw
        else:
            cur += ch; w += cw
    if cur or not out:
        out.append(cur)
    return out


# ---------- 低层图元（自由拼装用） ----------
def init_figure(width=6.4, height=8, title=""):
    """建画布返回 (fig, ax)。默认 6.4in 宽（贴近嵌入宽，缩放比≈1）。"""
    _ensure_font()
    fig, ax = plt.subplots(figsize=(width, height), dpi=200)
    ax.set_xlim(0, width)
    ax.set_ylim(0, height)
    ax.axis("off")
    if title:
        ax.set_title(title, fontsize=13, color=K, pad=10)
    return fig, ax


def box(ax, x, y, w, h, text, fs=12, lw=1.8):
    """左下角 (x,y) 的矩形框 + 居中文字（白底黑边）。"""
    ax.add_patch(Rectangle((x, y), w, h, facecolor="white", edgecolor=K, lw=lw, zorder=3))
    if text:
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                fontsize=fs, color=K, zorder=4)
    return (x + w / 2, y + h / 2)


def diamond(ax, cx, cy, w, h, text, fs=11):
    ax.add_patch(plt.Polygon([(cx, cy + h / 2), (cx + w / 2, cy), (cx, cy - h / 2), (cx - w / 2, cy)],
                             facecolor="white", edgecolor=K, lw=1.8, zorder=3))
    ax.text(cx, cy, text, ha="center", va="center", fontsize=fs, color=K, zorder=4)


def arrow(ax, p1, p2, text="", fs=10, head=True):
    """直箭头；带文字时文字置中点加白底（不压线）。head=False 为无头引线。"""
    style = "-|>" if head else "-"
    ax.add_patch(FancyArrowPatch(p1, p2, arrowstyle=style, mutation_scale=16,
                                 lw=1.6, color=K, shrinkA=1, shrinkB=1, zorder=2))
    if text:
        ax.text((p1[0] + p2[0]) / 2, (p1[1] + p2[1]) / 2, text, ha="center", va="center",
                fontsize=fs, color=K,
                bbox=dict(boxstyle="square,pad=0.12", fc="white", ec="none"), zorder=5)


def ftitle(ax, x, y, t, fs=13):
    ax.text(x, y, t, ha="center", fontsize=fs, color=K)


def save_figure(fig, filepath, dpi=200):
    fig.savefig(filepath, dpi=dpi, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"Saved: {filepath}")


# ---------- 高层：单列竖向自适应框图（流程 / 模块，首选） ----------
def _vstack(out_path, items, title="", fs=13, groups=None, feedback=None, down_labels=None):
    """单列竖向、框高随文字自适应、满版宽、线只在框间直走的框图核心。
    items: 每项 str 或 (标题, 详述)；groups: [(标签, 起idx, 止idx)] 虚线分组框；
    feedback: (从idx, 到idx, 标签) 左侧留白回流箭头；down_labels: {i: 标签} 标在第 i→i+1 箭头旁。"""
    _ensure_font()
    groups = groups or []
    down_labels = down_labels or {}
    W = 6.4
    has_fb = feedback is not None
    left = 1.5 if has_fb else 0.2
    right = 0.2
    box_x, box_w = left, W - left - right
    pad = 0.15
    line_h = fs * 1.5 / 72.0
    gap = 0.52
    inner_w = box_w - 2 * pad
    unit_in = fs / 72.0 / 2.0
    wrap = max(8, int(inner_w / unit_in * 0.80))

    laid = []
    for it in items:
        head, detail = (it if isinstance(it, (tuple, list)) else (it, None))
        hl = wrap_cjk(head, wrap)
        dl = wrap_cjk(detail, wrap) if detail else []
        h = (len(hl) + len(dl)) * line_h + 2 * pad
        laid.append((hl, dl, h))

    title_h = 0.46 if title else 0.12
    total_h = sum(b[2] for b in laid) + gap * (len(laid) - 1) + title_h + 0.18
    fig, ax = plt.subplots(figsize=(W, total_h), dpi=200)
    ax.set_xlim(0, W); ax.set_ylim(0, total_h); ax.axis("off")

    y = total_h - 0.09
    spans = []
    for (hl, dl, h) in laid:
        top, bot = y, y - h
        ax.add_patch(Rectangle((box_x, bot), box_w, h, facecolor="white", edgecolor=K, lw=1.8, zorder=3))
        ty = top - pad - line_h * 0.5
        for ln in hl:
            ax.text(box_x + box_w / 2, ty, ln, ha="center", va="center", fontsize=fs, color=K, zorder=4)
            ty -= line_h
        for ln in dl:
            ax.text(box_x + box_w / 2, ty, ln, ha="center", va="center", fontsize=fs - 1, color=K, zorder=4)
            ty -= line_h
        spans.append((top, bot))
        y = bot - gap

    cx = box_x + box_w / 2
    for i in range(len(spans) - 1):
        ax.add_patch(FancyArrowPatch((cx, spans[i][1]), (cx, spans[i + 1][0]),
                                     arrowstyle="-|>", mutation_scale=18, lw=1.8, color=K, zorder=2))
        if i in down_labels:
            ax.text(cx + 0.12, (spans[i][1] + spans[i + 1][0]) / 2, down_labels[i],
                    ha="left", va="center", fontsize=fs - 1, color=K,
                    bbox=dict(boxstyle="square,pad=0.1", fc="white", ec="none"), zorder=5)

    for (label, s, e) in groups:
        gt, gb = spans[s][0] + 0.30, spans[e][1] - 0.12
        gx0, gx1 = box_x - 0.14, box_x + box_w + 0.14
        ax.add_patch(Rectangle((gx0, gb), gx1 - gx0, gt - gb, fill=False, ec=K, lw=1.4,
                               ls=(0, (5, 4)), zorder=1))
        ax.text(gx0 + 0.1, gt - 0.03, label, ha="left", va="top", fontsize=fs - 1, color=K,
                bbox=dict(boxstyle="square,pad=0.12", fc="white", ec="none"), zorder=5)

    if has_fb:
        f, t, lab = feedback
        gx = left * 0.40
        yf = (spans[f][0] + spans[f][1]) / 2
        yt = (spans[t][0] + spans[t][1]) / 2
        ax.add_patch(FancyArrowPatch((box_x, yf), (gx, yf), arrowstyle="-", lw=1.7, color=K, zorder=2))
        ax.add_patch(FancyArrowPatch((gx, yf), (gx, yt), arrowstyle="-", lw=1.7, color=K, zorder=2))
        ax.add_patch(FancyArrowPatch((gx, yt), (box_x, yt), arrowstyle="-|>", mutation_scale=18,
                                     lw=1.7, color=K, zorder=2))
        for j, ln in enumerate(wrap_cjk(lab, max(4, int(left / unit_in * 0.8)))):
            ax.text(gx, (yf + yt) / 2 - (j - 0.5) * line_h, ln, ha="center", va="center",
                    fontsize=fs - 1, color=K,
                    bbox=dict(boxstyle="square,pad=0.1", fc="white", ec="none"), zorder=5)

    if title:
        ax.text(W / 2, 0.21, title, ha="center", va="center", fontsize=fs + 1, color=K, zorder=4)
    fig.savefig(out_path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"Saved: {out_path}")
    return out_path


def vflow(out_path, steps, title="", down_labels=None, fs=13):
    """竖向流程图（单列、框高自适应、满版宽、线不穿框）。steps 每项 str 或 (标题, 详述)。
    分支用 down_labels={步骤idx: '是'} 标注，别画右侧菱形旁注。"""
    return _vstack(out_path, steps, title=title, fs=fs, down_labels=down_labels)


def vmodules(out_path, mods, title="", groups=None, feedback=None, fs=13):
    """竖向模块框图（单列自适应）。mods 每项 str 或 (模块名, 职责)；
    groups=[(标签, 起, 止)] 虚线分组；feedback=(从idx, 到idx, 标签) 左侧回流箭头。"""
    return _vstack(out_path, mods, title=title, fs=fs, groups=groups, feedback=feedback)


# ---------- 扩展图型（同风格黑白自适应，覆盖五类专利） ----------
def draw_tree(out_path, tree, title="", fs=12):
    """树/决策分支图（回退链、决策树）。tree: {'text': 根, 'children': [...]}，
    子节点可带 'label' 标在连线上。自上而下自适应布局。"""
    _ensure_font()
    unit = fs / 72.0 / 2.0
    pad = 0.18
    line_h = fs * 1.5 / 72.0
    wrap_w = 16
    h_gap, v_gap = 0.45, 0.6

    def laid(t):
        lines = wrap_cjk(t["text"], wrap_w)
        w = max(1.6, max(disp_w(l) for l in lines) * unit + 2 * pad)
        h = len(lines) * line_h + 2 * pad
        return lines, w, h

    memo = {}

    def subtree_w(t):
        if id(t) in memo:
            return memo[id(t)]
        _, w, _ = laid(t)
        ch = t.get("children") or []
        r = max(w, (sum(subtree_w(c) for c in ch) + h_gap * (len(ch) - 1)) if ch else 0)
        memo[id(t)] = r
        return r

    def depth(t):
        ch = t.get("children") or []
        return 1 + max((depth(c) for c in ch), default=0)

    levels = {}

    def collect(t, d):
        levels.setdefault(d, []).append(t)
        for c in (t.get("children") or []):
            collect(c, d + 1)

    collect(tree, 0)
    D = depth(tree)
    level_h = {d: max(laid(t)[1] for t in ts) for d, ts in levels.items()}
    total_w = subtree_w(tree)
    W = max(6.4, total_w + 0.6)
    ys = {}
    yy = 0.2
    for d in range(D):
        ys[d] = yy
        yy += level_h[d]
        if d < D - 1:
            yy += v_gap
    H = yy + (0.4 if title else 0.15)

    fig, ax = plt.subplots(figsize=(W, H), dpi=200)
    ax.set_xlim(0, W); ax.set_ylim(0, H); ax.axis("off")
    pos = {}

    def place(t, x0, d):
        sw = subtree_w(t)
        _, w, h = laid(t)
        cx = x0 + sw / 2
        yb = H - 0.1 - ys[d] - h
        pos[id(t)] = (cx, yb, w, h)
        ch = t.get("children") or []
        if ch:
            chw = sum(subtree_w(c) for c in ch) + h_gap * (len(ch) - 1)
            cx0 = cx - chw / 2
            for c in ch:
                csw = subtree_w(c)
                place(c, cx0, d + 1)
                cx0 += csw + h_gap

    place(tree, (W - total_w) / 2, 0)

    def draw(t, d):
        lines, w, h = laid(t)
        cx, yb, _, _ = pos[id(t)]
        ax.add_patch(Rectangle((cx - w / 2, yb), w, h, facecolor="white", edgecolor=K, lw=1.8, zorder=3))
        ty = yb + h - pad - line_h * 0.5
        for ln in lines:
            ax.text(cx, ty, ln, ha="center", va="center", fontsize=fs, color=K, zorder=4)
            ty -= line_h
        for c in (t.get("children") or []):
            ccx, cyb, cw, chh = pos[id(c)]
            ax.add_patch(FancyArrowPatch((cx, yb), (ccx, cyb + chh), arrowstyle="-|>",
                                         mutation_scale=16, lw=1.6, color=K, zorder=2))
            if c.get("label"):
                ax.text((cx + ccx) / 2, (yb + cyb + chh) / 2, c["label"], ha="center",
                        va="center", fontsize=fs - 2, color=K,
                        bbox=dict(boxstyle="square,pad=0.1", fc="white", ec="none"), zorder=5)
            draw(c, d + 1)

    draw(tree, 0)
    if title:
        ax.text(W / 2, 0.14, title, ha="center", va="center", fontsize=fs + 1, color=K)
    fig.savefig(out_path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"Saved: {out_path}")
    return out_path


def draw_sequence(out_path, actors, messages, title="", fs=11):
    """时序图（模块间调用）。actors: 名称列表（左→右）；
    messages: [{'frm': i, 'to': j, 'text': 说明}]，to==frm 为自调用。"""
    _ensure_font()
    unit = fs / 72.0 / 2.0
    n = len(actors)
    W = max(6.4, 1.2 + n * 2.2)
    msg_gap = 0.55
    top = 0.8
    H = top + 1.0 + len(messages) * msg_gap + 0.9
    fig, ax = plt.subplots(figsize=(W, H), dpi=200)
    ax.set_xlim(0, W); ax.set_ylim(0, H); ax.axis("off")

    xs = [0.6 + (W - 1.2) * (i + 0.5) / n for i in range(n)]
    box_h = 0.6
    box_y = H - 0.2 - box_h
    for i, name in enumerate(actors):
        w = max(1.6, disp_w(name) * unit + 0.4)
        ax.add_patch(Rectangle((xs[i] - w / 2, box_y), w, box_h, facecolor="white",
                               edgecolor=K, lw=1.8, zorder=3))
        ax.text(xs[i], box_y + box_h / 2, name, ha="center", va="center",
                fontsize=fs, color=K, fontweight="bold", zorder=4)
    bottom = 0.7
    for x in xs:
        ax.plot([x, x], [box_y, bottom], color=K, ls=(0, (4, 3)), lw=1.0, zorder=1)

    y = box_y - 0.5
    for msg in messages:
        i, j = msg["frm"], msg["to"]
        if i == j:
            ax.add_patch(FancyArrowPatch((xs[i], y), (xs[i] + 0.7, y), arrowstyle="-", lw=1.6, color=K, zorder=2))
            ax.add_patch(FancyArrowPatch((xs[i] + 0.7, y), (xs[i] + 0.7, y - 0.3), arrowstyle="-", lw=1.6, color=K, zorder=2))
            ax.add_patch(FancyArrowPatch((xs[i] + 0.7, y - 0.3), (xs[i], y - 0.3), arrowstyle="-|>",
                                         mutation_scale=16, lw=1.6, color=K, zorder=2))
            ax.text(xs[i] + 0.8, y - 0.15, msg["text"], ha="left", va="center", fontsize=fs - 1,
                    color=K, bbox=dict(boxstyle="square,pad=0.1", fc="white", ec="none"), zorder=5)
        else:
            ax.add_patch(FancyArrowPatch((xs[i], y), (xs[j], y), arrowstyle="-|>",
                                         mutation_scale=16, lw=1.6, color=K, zorder=2))
            ax.text((xs[i] + xs[j]) / 2, y + 0.14, msg["text"], ha="center", va="bottom",
                    fontsize=fs - 1, color=K,
                    bbox=dict(boxstyle="square,pad=0.1", fc="white", ec="none"), zorder=5)
        y -= msg_gap

    if title:
        ax.text(W / 2, 0.22, title, ha="center", va="center", fontsize=fs + 1, color=K)
    fig.savefig(out_path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"Saved: {out_path}")
    return out_path


def draw_data_flow(out_path, nodes, flows, title="", fs=12):
    """数据流图。nodes: 有序 dict {名称: 说明或None}（自上而下单列）；
    flows: [{'frm': 名称, 'to': 名称, 'data': 数据说明}]，相邻向下走中线，其余走右侧 gutter。"""
    _ensure_font()
    names = list(nodes.keys())
    items = [(n, nodes[n]) if nodes[n] else n for n in names]
    idx = {n: i for i, n in enumerate(names)}
    # 复用 _vstack 布局前，先把非相邻 flows 转成右侧标注不便，直接自绘
    unit = fs / 72.0 / 2.0
    W = 6.4
    left, right = 0.2, 1.3
    box_x, box_w = left, W - left - right
    pad = 0.15
    line_h = fs * 1.5 / 72.0
    gap = 0.6
    wrap = max(8, int((box_w - 2 * pad) / unit * 0.80))

    laid = []
    for n in names:
        d = nodes[n]
        hl = wrap_cjk(n, wrap)
        dl = wrap_cjk(d, wrap) if d else []
        h = (len(hl) + len(dl)) * line_h + 2 * pad
        laid.append((hl, dl, h))
    title_h = 0.46 if title else 0.12
    total_h = sum(b[2] for b in laid) + gap * (len(laid) - 1) + title_h + 0.18
    fig, ax = plt.subplots(figsize=(W, total_h), dpi=200)
    ax.set_xlim(0, W); ax.set_ylim(0, total_h); ax.axis("off")

    y = total_h - 0.09
    spans = {}
    for n, (hl, dl, h) in zip(names, laid):
        top, bot = y, y - h
        ax.add_patch(Rectangle((box_x, bot), box_w, h, facecolor="white", edgecolor=K, lw=1.8, zorder=3))
        ty = top - pad - line_h * 0.5
        for ln in hl:
            ax.text(box_x + box_w / 2, ty, ln, ha="center", va="center", fontsize=fs, color=K, zorder=4)
            ty -= line_h
        for ln in dl:
            ax.text(box_x + box_w / 2, ty, ln, ha="center", va="center", fontsize=fs - 1, color=K, zorder=4)
            ty -= line_h
        spans[n] = (top, bot)
        y = bot - gap

    cx = box_x + box_w / 2
    gx = W - right * 0.45
    for f in flows:
        a, b = f["frm"], f["to"]
        ya = (spans[a][0] + spans[a][1]) / 2
        yb = (spans[b][0] + spans[b][1]) / 2
        if idx[b] == idx[a] + 1:
            ax.add_patch(FancyArrowPatch((cx, spans[a][1]), (cx, spans[b][0]), arrowstyle="-|>",
                                         mutation_scale=18, lw=1.8, color=K, zorder=2))
            ax.text(cx + 0.12, (spans[a][1] + spans[b][0]) / 2, f.get("data", ""),
                    ha="left", va="center", fontsize=fs - 1, color=K,
                    bbox=dict(boxstyle="square,pad=0.1", fc="white", ec="none"), zorder=5)
        else:
            x0 = box_x + box_w
            ax.add_patch(FancyArrowPatch((x0, ya), (gx, ya), arrowstyle="-", lw=1.6, color=K, zorder=2))
            ax.add_patch(FancyArrowPatch((gx, ya), (gx, yb), arrowstyle="-", lw=1.6, color=K, zorder=2))
            ax.add_patch(FancyArrowPatch((gx, yb), (x0, yb), arrowstyle="-|>", mutation_scale=18,
                                         lw=1.6, color=K, zorder=2))
            for j, ln in enumerate(wrap_cjk(f.get("data", ""), 10)):
                ax.text(gx + 0.08, (ya + yb) / 2 - (j - 0.5) * line_h, ln, ha="left", va="center",
                        fontsize=fs - 1, color=K,
                        bbox=dict(boxstyle="square,pad=0.1", fc="white", ec="none"), zorder=5)

    if title:
        ax.text(W / 2, 0.21, title, ha="center", va="center", fontsize=fs + 1, color=K, zorder=4)
    fig.savefig(out_path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"Saved: {out_path}")
    return out_path


def draw_component_callout(out_path, parts, title="", body="整体结构", fs=12):
    """部件标注图（机械结构类）：中心主体框 + 引线 + 黑白编号圈。
    parts: [{'num': '1', 'name': 部件名, 'lx': x, 'ly': y}]，lx/ly 为 6.4×6 画布坐标。"""
    _ensure_font()
    W, H = 6.4, 6.0
    fig, ax = plt.subplots(figsize=(W, H), dpi=200)
    ax.set_xlim(0, W); ax.set_ylim(0, H); ax.axis("off")
    cx, cy = W / 2, H / 2
    bw, bh = 2.2, 1.2
    ax.add_patch(Rectangle((cx - bw / 2, cy - bh / 2), bw, bh, facecolor="white",
                           edgecolor=K, lw=1.8, zorder=3))
    ax.text(cx, cy, body, ha="center", va="center", fontsize=fs, color=K, fontweight="bold", zorder=4)
    for p in parts:
        lx, ly = p["lx"], p["ly"]
        # 引线从主体边缘起，避免穿过主体
        dx, dy = lx - cx, ly - cy
        dist = math.hypot(dx, dy) or 1
        sx, sy = cx + dx / dist * (bw / 2 * 0.9), cy + dy / dist * (bh / 2 * 0.9)
        ax.add_patch(FancyArrowPatch((sx, sy), (lx, ly), arrowstyle="-", lw=1.2, color=K, zorder=2))
        ax.add_patch(plt.Circle((lx, ly), 0.22, facecolor="white", edgecolor=K, lw=1.4, zorder=5))
        ax.text(lx, ly, p["num"], ha="center", va="center", fontsize=fs - 2, color=K,
                fontweight="bold", zorder=6)
        # 名字放在圆圈外侧（左半画布向左排，右半向右排），不压引线
        if lx <= W / 2:
            ax.text(lx - 0.32, ly, p["name"], ha="right", va="center", fontsize=fs - 1, color=K, zorder=6)
        else:
            ax.text(lx + 0.32, ly, p["name"], ha="left", va="center", fontsize=fs - 1, color=K, zorder=6)
    if title:
        ax.text(W / 2, 0.2, title, ha="center", va="center", fontsize=fs + 1, color=K)
    fig.savefig(out_path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"Saved: {out_path}")
    return out_path


def draw_six_views(out_path, labels=None, title="", fs=12):
    """外观设计六面视图占位框（2 行×3 列），供贴产品图或手绘。"""
    _ensure_font()
    names = {"main": "主视图", "back": "后视图", "left": "左视图",
             "right": "右视图", "top": "俯视图", "bottom": "仰视图"}
    if labels:
        names.update(labels)
    order = [["main", "back", "left"], ["right", "top", "bottom"]]
    W, H = 6.4, 5.4
    fig, ax = plt.subplots(figsize=(W, H), dpi=200)
    ax.set_xlim(0, W); ax.set_ylim(0, H); ax.axis("off")
    bw, bh = 1.7, 1.5
    for r, rowk in enumerate(order):
        for c, key in enumerate(rowk):
            x = W / 6 * (2 * c + 1)
            y = H - (H / 4 * (2 * r + 1)) + 0.35
            ax.add_patch(Rectangle((x - bw / 2, y - bh / 2), bw, bh, facecolor="white",
                                   edgecolor=K, lw=1.5, zorder=3))
            ax.text(x, y - bh / 2 - 0.22, names[key], ha="center", va="top",
                    fontsize=fs - 1, color=K, fontweight="bold")
    if title:
        ax.text(W / 2, 0.2, title, ha="center", va="center", fontsize=fs + 1, color=K)
    fig.savefig(out_path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"Saved: {out_path}")
    return out_path


def json_examples(out_path, pos_title, pos_text, neg_title, neg_text, title=""):
    """上下堆叠两个 JSON 文本块（黑框自适应包裹文字）；含中文必须用 CJK 字体。
    仅当发明确有结构化输出值得展示时才用。"""
    _ensure_font()
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(6.8, 9.6))
    for ax, ttl, txt in [(ax1, pos_title, pos_text), (ax2, neg_title, neg_text)]:
        ax.axis("off")
        ax.set_title(ttl, fontsize=13, color=K, pad=8)
        ax.text(0.5, 0.5, txt, transform=ax.transAxes, fontsize=12, ha="center", va="center",
                color=K, fontfamily=FONT, linespacing=1.6,
                bbox=dict(boxstyle="square,pad=0.8", facecolor="white", edgecolor=K, linewidth=1.4))
    if title:
        fig.suptitle(title, fontsize=13.5, color=K, y=0.015)
    fig.subplots_adjust(hspace=0.32, top=0.95, bottom=0.05)
    fig.savefig(out_path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"Saved: {out_path}")
    return out_path


def insert_images_to_cell(doc_path, table_idx, row_idx, col_idx, fig_map, width_cm=None):
    """将图片插入 Word 指定表格单元格：找到含标记文字（如'附图1'）的段落后插入居中图片。

    width_cm 默认 None = 按模板该列实际宽度自动取（列宽 - 0.3cm 余量），保证不越界；
    显式传值时也会夹紧到列宽以内（官方模板内容列约 12.98cm，传 15 必越界）。
    返回 {'inserted': [标记...], 'missing': [标记...]}。
    missing 非空表示正文缺该（附图N）标记——调用方必须报错处理，不得静默交付。
    """
    from docx import Document
    from docx.shared import Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    doc = Document(doc_path)
    table = doc.tables[table_idx]
    cell = table.cell(row_idx, col_idx)
    tc = cell._tc

    # 由 tblGrid 计算单元格实际宽度（twips→cm），处理 gridSpan 合并列
    grid = [int(g.get(qn('w:w'))) for g in table._tbl.tblGrid.findall(qn('w:gridCol'))]
    span_elem = tc.tcPr.find(qn('w:gridSpan')) if tc.tcPr is not None else None
    span = int(span_elem.get(qn('w:val'))) if span_elem is not None else 1
    cell_w_cm = sum(grid[col_idx:col_idx + span]) / 1440 * 2.54
    max_w = cell_w_cm - 0.3
    eff_w = min(width_cm, max_w) if width_cm else max_w

    inserted, missing = [], []

    for marker, img_path in fig_map.items():
        found = False
        for p_elem in list(tc.iterchildren(qn('w:p'))):
            p_text = ''.join(node.text or '' for node in p_elem.iter(qn('w:t')))
            if marker in p_text:
                new_p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                jc = OxmlElement('w:jc')
                jc.set(qn('w:val'), 'center')
                pPr.append(jc)
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:before'), '120')
                spacing.set(qn('w:after'), '120')
                pPr.append(spacing)
                new_p.append(pPr)
                p_elem.addnext(new_p)

                new_para = Paragraph(new_p, cell)
                run = new_para.add_run()
                run.add_picture(img_path, width=Cm(eff_w))
                print(f"Inserted {marker} -> {img_path} (width {eff_w:.2f}cm, 列宽 {cell_w_cm:.2f}cm)")
                inserted.append(marker)
                found = True
                break
        if not found:
            print(f"⚠️ MARKER NOT FOUND: {marker}（3.2 单元格无「{marker}」标记，图未嵌入）")
            missing.append(marker)

    doc.save(doc_path)
    print(f"Document saved: {doc_path}")
    return {"inserted": inserted, "missing": missing}
