"""
专利技术交底书保密/脱敏自检脚本。

扫描生成的 .docx，命中疑似敏感信息即告警（仅提示，不自动删除，交人工复核）。
是 SKILL.md Step 6 保密自检的执行工具。

依赖：pip install python-docx
"""
import re
from docx import Document


# === 检测规则 ===
# 每条规则: (类别, 正则, 说明, 是否忽略联系人表)
_RULES = [
    ('内网域名', re.compile(r'[\w.-]+\.sf-express\.com'),
     '疑似顺丰内部系统域名', True),
    ('内网IP', re.compile(r'\b(?:10|172|192)\.\d{1,3}\.\d{1,3}\.\d{1,3}\b'),
     '疑似内网 IP 地址', False),
    ('手机号', re.compile(r'(?<!\d)1[3-9]\d{9}(?!\d)'),
     '疑似手机号（联系人表除外需脱敏）', True),
    ('工号', re.compile(r'(?<!\d)0\d{7}(?!\d)'),
     '疑似员工工号', True),
    ('邮箱', re.compile(r'[\w.+-]+@[\w-]+\.[\w.-]+'),
     '疑似邮箱（联系人表除外需脱敏）', True),
]

# 疑似源代码：出现这些编程符号/关键字的密集行
_CODE_KEYWORDS = re.compile(
    r'\b(def|function|class|import|return|public|private|void|const|let|var)\b'
    r'|[{};]\s*$|=>|->|::'
)

# 未解释英文缩写：连续 2+ 大写字母，且未跟中文括号解释
_ABBR = re.compile(r'(?<![A-Za-z])[A-Z]{2,}(?![A-Za-z])')

# 疑似未脱敏实验数据：一行内出现 4 个以上精确数字（含小数）
_DENSE_NUM = re.compile(r'(?:(?<!\w)\d+(?:\.\d+)?(?!\w).*?){4,}')


def _iter_text_units(doc):
    """产出 (定位描述, 文本, 是否属于联系人表) 的迭代器。"""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            yield (f'正文段落#{i}', p.text, False)
    for ti, t in enumerate(doc.tables):
        is_contact = (ti == 0)  # 表0 = 联系人信息表
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                txt = cell.text.strip()
                if txt:
                    yield (f'表{ti}[{ri},{ci}]', txt, is_contact)


def scan(doc_path):
    """
    扫描文档，返回告警列表。每项: dict(category, location, snippet, note)。
    同时打印可读报告。
    """
    doc = Document(doc_path)
    issues = []

    for loc, text, is_contact in _iter_text_units(doc):
        # 规则类检测
        for cat, pat, note, skip_contact in _RULES:
            if skip_contact and is_contact:
                continue
            for m in pat.finditer(text):
                issues.append({'category': cat, 'location': loc,
                               'snippet': m.group(0), 'note': note})

        # 源代码检测（按行）
        for line in text.split('\n'):
            if len(_CODE_KEYWORDS.findall(line)) >= 2:
                issues.append({'category': '疑似源代码', 'location': loc,
                               'snippet': line.strip()[:60],
                               'note': '疑似源代码片段，专利交底书不应含未公开源码'})
                break

        # 未解释缩写
        for m in set(_ABBR.findall(text)):
            # 若缩写后紧跟中文括号解释则认为已解释
            if re.search(re.escape(m) + r'\s*[（(]', text):
                continue
            issues.append({'category': '未解释缩写', 'location': loc,
                           'snippet': m, 'note': '英文缩写建议给出全称+中文解释'})

        # 密集数字（仅正文，避免误报表格里的编号）
        if not is_contact and _DENSE_NUM.search(text):
            # 逐行判断，排除以子节编号(如 2.1 / 3.2.1)开头的引导行
            for line in text.split('\n'):
                s = line.strip()
                if re.match(r'^\d+(\.\d+)+', s):
                    continue
                nums = re.findall(r'(?<!\w)\d+(?:\.\d+)?(?!\w)', s)
                if len(nums) >= 4:
                    issues.append({'category': '疑似未脱敏数据', 'location': loc,
                                   'snippet': ('、'.join(nums[:6]) + '…'),
                                   'note': '成片精确数字，确认是否为未脱敏实验数据'})
                    break

    _report(doc_path, issues)
    return issues


def _report(doc_path, issues):
    print(f'\n=== 保密/脱敏自检：{doc_path} ===')
    if not issues:
        print('✅ 未发现明显敏感信息（仍建议人工终审）')
        return
    # 按类别聚合
    by_cat = {}
    for it in issues:
        by_cat.setdefault(it['category'], []).append(it)
    print(f'⚠️  共 {len(issues)} 处待人工复核：\n')
    for cat, items in by_cat.items():
        print(f'【{cat}】{items[0]["note"]}  ({len(items)}处)')
        for it in items[:8]:
            print(f'   - {it["location"]}: {it["snippet"]}')
        if len(items) > 8:
            print(f'   … 另有 {len(items)-8} 处')
        print()
    print('注：本工具仅告警，不自动删除。请逐条人工确认后脱敏。')


if __name__ == '__main__':
    import sys
    if len(sys.argv) > 1:
        scan(sys.argv[1])
    else:
        print('用法: python compliance_check.py <文档.docx>')
