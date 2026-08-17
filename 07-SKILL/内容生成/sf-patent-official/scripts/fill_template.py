"""
专利技术交底书模板填充脚本。

把易错的 python-docx 手工定位表格、清除示例、填内容封装成一次调用。
所有 5 类模板共享相同的两表结构（由 fill_template 自动适配）：
  - 表0 联系人信息：3行4列，值位于 (1,1)姓名 (1,3)电话 (2,1)部门 (2,3)邮箱
  - 表1 详细技术信息：6行2列，5大章节内容位于第1列 row1~row5

依赖：pip install python-docx
"""
from docx import Document
from docx.oxml.ns import qn


# 表1 中「章节标题 -> 行号」映射（第0行是表头「详细技术信息」）
_SECTION_ROW = {
    'title': 1,           # 1、提案名称
    'background': 2,      # 2、背景技术
    'solution': 3,        # 3、详细方案
    'supplement': 4,      # 4、技术补充
    'statement': 5,       # 5、技术声明
}

# 5.1 实施情况四选一的标准表述
_IMPLEMENTATION = {
    'A': '□ 属于技术研发成果，未应用到实际产品',
    'B': '该方案已应用到实际产品，产品在公司内部应用',
    'C': '该方案确定应用到实际产品，将要或已经对外公开应用',
    'D': '不属于技术研发成果，未应用到实际产品',
}


def _set_cell_text(cell, text):
    """清空单元格全部段落后写入 text（多行按 \\n 拆分为多段）。保留首段样式。"""
    paras = cell.paragraphs
    # 删除除第一段外的所有段落
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    # 清空第一段的所有 run
    for r in list(first.runs):
        r._element.getparent().remove(r._element)
    lines = str(text).split('\n')
    first.add_run(lines[0])
    for line in lines[1:]:
        p = cell.add_paragraph()
        p.add_run(line)


def _compose_background(s):
    return (
        f"2.1 现有技术介绍\n{s.get('bg_intro', '')}\n\n"
        f"2.2 现有技术的缺点 / 本发明解决的技术问题\n{s.get('bg_problem', '')}"
    )


def _compose_solution(s):
    return (
        f"3.1 核心发明点\n{s.get('invent_points', '')}\n\n"
        f"3.2 完整技术方案描述\n{s.get('solution', '')}\n\n"
        f"3.3 技术效果\n{s.get('effects', '')}"
    )


def _compose_supplement(s):
    return (
        f"4.1 替代方案\n{s.get('alternatives', '')}\n\n"
        f"4.2 术语说明\n{s.get('terms', '')}\n\n"
        f"4.3 参考文献\n{s.get('refs', '')}"
    )


def _compose_statement(s):
    impl = _IMPLEMENTATION.get(s.get('implementation', 'A'), s.get('implementation', ''))
    return (
        f"5.1 实施情况\n{impl}\n\n"
        f"5.2 公开情况\n{s.get('disclosure', '')}"
    )


def fill_patent_doc(template, output, contact, sections):
    """
    填充专利技术交底书。

    template: 模板 .docx 路径（assets/ 下某一类）
    output:   输出 .docx 路径
    contact:  dict, 键 name / phone / dept / email
    sections: dict, 键见下。多子节章节由脚本自动拼装子节标题：
        title          -> 1、提案名称
        bg_intro       -> 2.1 现有技术介绍
        bg_problem     -> 2.2 现有技术缺点
        invent_points  -> 3.1 核心发明点
        solution       -> 3.2 完整技术方案（含「附图N」标记）
        effects        -> 3.3 技术效果
        alternatives   -> 4.1 替代方案
        terms          -> 4.2 术语说明
        refs           -> 4.3 参考文献
        implementation -> 5.1 实施情况，取值 'A'/'B'/'C'/'D' 或自定义文本
        disclosure     -> 5.2 公开情况
    """
    doc = Document(template)

    # --- 表0 联系人信息 ---
    t0 = doc.tables[0]
    _set_cell_text(t0.cell(1, 1), contact.get('name', ''))
    _set_cell_text(t0.cell(1, 3), contact.get('phone', ''))
    _set_cell_text(t0.cell(2, 1), contact.get('dept', ''))
    _set_cell_text(t0.cell(2, 3), contact.get('email', ''))

    # --- 表1 详细技术信息 ---
    t1 = doc.tables[1]
    _set_cell_text(t1.cell(_SECTION_ROW['title'], 1), sections.get('title', ''))
    _set_cell_text(t1.cell(_SECTION_ROW['background'], 1), _compose_background(sections))
    _set_cell_text(t1.cell(_SECTION_ROW['solution'], 1), _compose_solution(sections))
    _set_cell_text(t1.cell(_SECTION_ROW['supplement'], 1), _compose_supplement(sections))
    _set_cell_text(t1.cell(_SECTION_ROW['statement'], 1), _compose_statement(sections))

    doc.save(output)
    print(f"Filled document saved: {output}")
    print("下一步：调用 embed_figures(output, fig_map) 将附图嵌入 3.2 单元格（强制校验，缺标记即报错）")
    return output


def embed_figures(doc_path, fig_map, width_cm=None):
    """
    装配后必跑：把附图嵌入「表1 详细技术信息 / 第3行 3.2 详细方案 / 第1列」单元格。

    fig_map: {'附图1': 'figures/fig1_flow.png', ...}，键须与 3.2 正文（附图N）标记一致。
    强制校验：图片文件必须存在；每个标记必须在单元格中找到。任一不满足抛 RuntimeError——
    宁可中断流程，不得交付无附图的交底书。
    """
    import os
    for marker, path in fig_map.items():
        if not os.path.exists(path):
            raise RuntimeError(f"附图文件不存在：{path}（{marker}）。先回阶段 9 绘图。")
    if not fig_map:
        raise RuntimeError("fig_map 为空——交底书必须含 ≥2 张附图，回阶段 9。")

    import draw_utils
    report = draw_utils.insert_images_to_cell(
        doc_path, table_idx=1, row_idx=3, col_idx=1, fig_map=fig_map, width_cm=width_cm)
    if report["missing"]:
        raise RuntimeError(
            f"以下附图标记未能在 3.2 单元格中找到，图未嵌入：{report['missing']}。"
            f"检查正文是否含（附图N：…）标记且与 fig_map 键一致；回阶段 8/9 修正后重装配。")
    if len(report["inserted"]) < 2:
        raise RuntimeError(f"嵌入附图 {len(report['inserted'])} 张 < 2，不满足交底书最低要求。")
    print(f"✅ 附图嵌入校验通过：{len(report['inserted'])}/{len(fig_map)} 张")
    return report


if __name__ == '__main__':
    # 自测：用应用软件类模板生成一份样例
    import os
    here = os.path.dirname(os.path.abspath(__file__))
    tmpl = os.path.join(here, '..', 'assets', '应用软件类模板.docx')
    fill_patent_doc(
        template=tmpl,
        output=os.path.join(here, '_selftest_output.docx'),
        contact={'name': '张三', 'phone': '138xxxx', 'dept': '顺丰科技/AI平台', 'email': 'zs@sf-express.com'},
        sections={
            'title': '一种基于XXX的YYY方法',
            'bg_intro': '现有技术介绍……', 'bg_problem': '现有缺点……',
            'invent_points': '① …… ② ……', 'solution': '首先……（附图1：系统架构）',
            'effects': '效果1：提升10倍……',
            'alternatives': '替代方案1……', 'terms': 'TTFT：首字延迟', 'refs': '无',
            'implementation': 'B', 'disclosure': '暂无对外公开计划',
        },
    )
