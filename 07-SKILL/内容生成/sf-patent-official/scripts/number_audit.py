#!/usr/bin/env python3
"""数字对照脚本 — 阶段 13 事实锁定/反幻觉核验

目的：给终稿里每个量化数字找"出生证明"，无出处的标黄待核。

输入：① 成稿（md 或 docx，路径或原文字符串）
      ② 白名单：阶段 1 摘录 + 用户对话确认的数字，格式
         [{"value": "40%", "source": "《XX详设.md》第3节", "context": "分拣效率提升约40%"}, ...]
         value 支持区间（"2-4"、"2~4"、"2至4"），区间内数字视为有出处
输出：报告 dict + 标注稿（仅 docx 输入）：
- matched：命中白名单的数字（含等价换算：百分之三十=30%=0.3）
- unmatched：无出处数字列表（值 + 上下文）
- annotated：docx 输入时 unmatched 数字所在 run 加黄色高亮，另存为 *_标注.docx

实现要点：
- 抽取：阿拉伯数字+单位（%|ms|s|倍|万|亿|元|℃|GB|QPS|单/小时…）；区间（a-b、a~b）；
  中文表述（百分之X、X成、X倍）归一化后比数值
- 判定是确定性的（在/不在白名单），不做模糊放行

依赖：python-docx（仅 docx 输入时需要）
"""

import json
import os
import re
import sys

CN_DIG = {"零": 0, "一": 1, "二": 2, "两": 2, "三": 3, "四": 4,
          "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
_UNIT = r"(?:%|％|ms|s|倍|万|亿|元|℃|GB|TB|MB|QPS|单/小时|分钟|小时|天|个月|月|年|次)"
PATTERNS = [
    re.compile(r"\d+(?:\.\d+)?\s*[-~至]\s*\d+(?:\.\d+)?"),  # 区间整体视为一个数字串
    re.compile(r"\d+(?:\.\d+)?\s*" + _UNIT + "?"),
    re.compile(r"百分之[零一二两三四五六七八九十百]+"),
    re.compile(r"[零一二两三四五六七八九十]+成"),
    re.compile(r"[零一二两三四五六七八九十]+倍"),
]


def _cn_int(s: str) -> int:
    """简易中文整数（0–999）：三十 / 十五 / 一百二十 / 三"""
    if s in CN_DIG:
        return CN_DIG[s]
    total, num = 0, 0
    for ch in s:
        if ch in CN_DIG:
            num = CN_DIG[ch]
        elif ch == "十":
            total += (num or 1) * 10
            num = 0
        elif ch == "百":
            total += (num or 1) * 100
            num = 0
    return total + num


def _canonical(tok: str):
    """归一化为 (kind, 数值)；kind=percent 时数值为百分数（30% → 30）。"""
    m = re.match(r"百分之([零一二两三四五六七八九十百]+)", tok)
    if m:
        return ("percent", float(_cn_int(m.group(1))))
    m = re.match(r"([零一二两三四五六七八九十]+)成", tok)
    if m:
        return ("percent", float(_cn_int(m.group(1))) * 10)
    m = re.match(r"([零一二两三四五六七八九十]+)倍", tok)
    if m:
        return ("plain", float(_cn_int(m.group(1))))
    num = float(re.match(r"\d+(?:\.\d+)?", tok).group())
    if "%" in tok or "％" in tok:
        return ("percent", num)
    return ("plain", num)


def _wl_entry(value: str):
    """白名单 value → (kind, lo, hi)；区间返回闭区间，普通值 lo==hi。"""
    v = value.replace(" ", "")
    m = re.match(r"^(\d+(?:\.\d+)?)(%|％)?[-~至](\d+(?:\.\d+)?)(%|％)?" + _UNIT + r"?$", v)
    if m:
        lo = _canonical(m.group(1) + (m.group(2) or ""))
        hi = _canonical(m.group(3) + (m.group(4) or ""))
        return (lo[0], min(lo[1], hi[1]), max(lo[1], hi[1]))
    first = re.match(
        r"\d+(?:\.\d+)?\s*(?:%|％)?|百分之[零一二两三四五六七八九十百]+|"
        r"[零一二两三四五六七八九十]+[成倍]", v)
    c = _canonical(first.group() if first else v)
    return (c[0], c[1], c[1])


def _match(canon, entry) -> bool:
    kind, lo, hi = entry
    v = canon[1]
    if canon[0] != kind:  # 百分之三十 = 30% = 0.3 等价换算
        if canon[0] == "percent" and kind == "plain":
            v = v / 100
        elif canon[0] == "plain" and kind == "percent" and v <= 1:
            v = v * 100
        else:
            return False
    return lo - 1e-9 <= v <= hi + 1e-9


def extract_numbers(text: str) -> list:
    hits = []
    for pat in PATTERNS:
        for m in pat.finditer(text):
            hits.append((m.start(), m.end(), m.group().replace(" ", "")))
    hits.sort()
    out, last_end = [], -1
    for s, e, tok in hits:  # 区间模式优先，去掉重叠的阿拉伯/中文碎片
        if s < last_end:
            continue
        out.append((s, tok))
        last_end = e
    return out


def _context(text: str, pos: int, width: int = 20) -> str:
    s, e = max(0, pos - width), min(len(text), pos + width)
    return text[s:e].replace("\n", " ")


def _read_docx(path: str):
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(c.text for c in row.cells))
    return doc, "\n".join(parts)


def audit(text_or_path: str, whitelist: list) -> dict:
    doc = None
    if isinstance(text_or_path, str) and os.path.isfile(text_or_path) \
            and text_or_path.lower().endswith(".docx"):
        doc, text = _read_docx(text_or_path)
    elif isinstance(text_or_path, str) and os.path.isfile(text_or_path):
        with open(text_or_path, encoding="utf-8") as f:
            text = f.read()
    else:
        text = text_or_path

    entries = [(_wl_entry(w["value"]), w) for w in whitelist]

    matched, unmatched = [], []
    for pos, tok in extract_numbers(text):
        canon = _canonical(tok)
        hit = next((w for e, w in entries if _match(canon, e)), None)
        item = {"value": tok, "context": _context(text, pos)}
        if hit:
            matched.append({**item, "source": hit.get("source", "")})
        else:
            unmatched.append(item)

    report = {"matched": matched, "unmatched": unmatched}

    if doc is not None and unmatched:
        from docx.enum.text import WD_COLOR_INDEX
        for para in doc.paragraphs:
            for run in para.runs:
                if any(u["value"] in run.text.replace(" ", "") for u in unmatched):
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        out_path = text_or_path[:-5] + "_标注.docx"
        doc.save(out_path)
        report["annotated"] = out_path
    return report


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法：python number_audit.py <成稿.md|.docx> <白名单.json>")
        sys.exit(1)
    with open(sys.argv[2], encoding="utf-8") as f:
        wl = json.load(f)
    print(json.dumps(audit(sys.argv[1], wl), ensure_ascii=False, indent=2))
