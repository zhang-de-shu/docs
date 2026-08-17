#!/usr/bin/env python3
"""术语一致性扫描脚本 — 阶段 8 正文撰写（术语锁）

目的：检测全文术语漂移（同一对象多种叫法），防审查员判"不清楚"。

输入：① 术语表（dict：标准名 → 别名列表，阶段 8 用户确认过）
      ② 成稿文本（md 或 docx，路径或原文字符串）
输出：报告 dict：
- term_counts：各标准名出现次数
- alias_residuals：别名残留（标准名/别名/次数/上下文），应为 0
- candidates：高频名词（≥2 字、出现 ≥3 次）中不在术语表的候选词，提示人工判断

实现要点：
- 纯规则：计数 + 子串匹配，不需要模型
- docx 输入用 python-docx 抽段落与表格文本
- 候选词用重复 n-gram 挖掘（2–6 字），最长优先去重（子串与超串同频则弃子串）

依赖：python-docx（仅 docx 输入时需要）
"""

import json
import os
import re
import sys
from collections import Counter

CN_RUN = re.compile(r'[一-鿿]+')


def load_text(path_or_text: str) -> str:
    if isinstance(path_or_text, str) and os.path.isfile(path_or_text):
        if path_or_text.lower().endswith(".docx"):
            from docx import Document
            doc = Document(path_or_text)
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" ".join(c.text for c in row.cells))
            return "\n".join(parts)
        with open(path_or_text, encoding="utf-8") as f:
            return f.read()
    return path_or_text


def _context(text: str, pos: int, width: int = 20) -> str:
    s, e = max(0, pos - width), min(len(text), pos + width)
    return text[s:e].replace("\n", " ")


def _mine_candidates(text: str, known: set, min_len=2, max_len=6, min_count=3, top=50) -> list:
    cnt = Counter()
    for run in CN_RUN.findall(text):
        for n in range(min_len, min(max_len, len(run)) + 1):
            for i in range(0, len(run) - n + 1):
                cnt[run[i:i + n]] += 1
    freq = {w: c for w, c in cnt.items() if c >= min_count}
    kept = []
    for w in sorted(freq, key=lambda w: (-len(w), w)):  # 最长优先
        if any(w in k and freq[k] == freq[w] for k in kept):
            continue  # 与超串同频，属同一词的碎片
        kept.append(w)
    def _is_fragment(w: str) -> bool:
        """w 是已知词的子串或循环移位碎片（如 缓存系统缓存系统 → 存系统缓存系）则弃。"""
        for k in known:
            if w == k:
                return True
            if len(w) <= len(k) and w in k:
                return True
            if len(w) > len(k) and w in k * (len(w) // len(k) + 2):
                return True
        return False

    cands = [w for w in kept if not _is_fragment(w)]
    cands.sort(key=lambda w: -freq[w])
    return [{"word": w, "count": freq[w]} for w in cands[:top]]


def scan(terms: dict, text: str) -> dict:
    text = load_text(text)
    known = set(terms)
    for aliases in terms.values():
        known.update(aliases)

    term_counts = {name: text.count(name) for name in terms}

    alias_residuals = []
    for name, aliases in terms.items():
        for alias in aliases:
            contexts, pos = [], 0
            while True:
                pos = text.find(alias, pos)
                if pos < 0:
                    break
                contexts.append(_context(text, pos))
                pos += len(alias)
            if contexts:
                alias_residuals.append({
                    "standard": name, "alias": alias,
                    "count": len(contexts), "contexts": contexts[:5],
                })

    return {
        "term_counts": term_counts,
        "alias_residuals": alias_residuals,
        "candidates": _mine_candidates(text, known),
    }


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法：python term_consistency.py <术语表.json> <成稿.md|.docx|文本>")
        sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as f:
        terms = json.load(f)
    print(json.dumps(scan(terms, sys.argv[2]), ensure_ascii=False, indent=2))
