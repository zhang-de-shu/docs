"""
专利技术交底书附图绘制工具库。
使用 matplotlib 生成专利风格的流程图、架构图、机制示意图。
依赖：pip install matplotlib
"""
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
import os

# === 专利附图调色板 ===
COLORS = {
    'primary': '#2B5797',
    'secondary': '#4A90D9',
    'accent': '#E8792B',
    'success': '#2D8C3C',
    'danger': '#D32F2F',
    'purple': '#7B1FA2',
    'light_bg': '#F0F4FA',
    'light_blue': '#D6E4F0',
    'light_orange': '#FDE8D0',
    'light_green': '#D5EDDA',
    'light_purple': '#F3E5F5',
    'light_yellow': '#FFF9C4',
    'white': '#FFFFFF',
    'gray': '#E8E8E8',
    'dark': '#333333',
}


def init_figure(width=14, height=8, title='', xlim=None, ylim=None):
    """创建画布，返回 (fig, ax)。自动设置中文字体和隐藏坐标轴。"""
    plt.rcParams['font.family'] = ['Songti SC', 'STSong', 'SimSun', 'sans-serif']
    plt.rcParams['axes.unicode_minus'] = False

    fig, ax = plt.subplots(1, 1, figsize=(width, height))
    if xlim:
        ax.set_xlim(*xlim)
    else:
        ax.set_xlim(0, width)
    if ylim:
        ax.set_ylim(*ylim)
    else:
        ax.set_ylim(0, height)
    ax.axis('off')
    if title:
        ax.set_title(title, fontsize=14, fontweight='bold', pad=15, color=COLORS['dark'])
    return fig, ax


def draw_box(ax, x, y, w, h, text='', facecolor='#FFFFFF', edgecolor='#999999',
             fontsize=10, fontcolor='#333333', fontweight='normal', lw=1.2, pad=0.02):
    """在 (x,y) 为中心绘制圆角矩形，可包含居中文字。"""
    box = FancyBboxPatch((x - w/2, y - h/2), w, h,
                         boxstyle=f"round,pad={pad}", facecolor=facecolor,
                         edgecolor=edgecolor, linewidth=lw)
    ax.add_patch(box)
    if text:
        ax.text(x, y, text, ha='center', va='center', fontsize=fontsize,
                color=fontcolor, fontweight=fontweight, wrap=True)
    return box


def draw_arrow(ax, x1, y1, x2, y2, color='#666666', lw=1.5, style='->'):
    """从 (x1,y1) 到 (x2,y2) 绘制箭头。"""
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle=style, color=color, lw=lw))


def draw_circle_label(ax, x, y, text, radius=0.22, facecolor='#4A90D9',
                      fontsize=7, fontcolor='white'):
    """绘制带文字的小圆圈标签（用于步骤编号等）。"""
    circle = plt.Circle((x, y), radius, facecolor=facecolor,
                        edgecolor='white', linewidth=1.5, zorder=5)
    ax.add_patch(circle)
    ax.text(x, y, text, ha='center', va='center', fontsize=fontsize,
            color=fontcolor, fontweight='bold', zorder=6)


def draw_layer_bg(ax, x, y, w, h, title, bg_color, edge_color, alpha=0.5):
    """绘制一个带标题的层级背景区域（用于架构图分层）。"""
    rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05",
                          facecolor=bg_color, edgecolor=edge_color,
                          linewidth=1.5, alpha=alpha)
    ax.add_patch(rect)
    ax.text(x + 0.3, y + h - 0.25, title, fontsize=10, fontweight='bold',
            color=edge_color, va='top')


def draw_badge(ax, x, y, text, bg_color='#FFEBEE', edge_color='#D32F2F',
               fontsize=8, fontcolor=None):
    """绘制一个小型标注徽章。"""
    if fontcolor is None:
        fontcolor = edge_color
    ax.text(x, y, text, ha='center', va='center', fontsize=fontsize,
            color=fontcolor,
            bbox=dict(boxstyle='round,pad=0.2', facecolor=bg_color,
                     edgecolor=edge_color, lw=0.8))


def save_figure(fig, filepath, dpi=200):
    """保存图片并关闭画布。"""
    fig.savefig(filepath, dpi=dpi, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    print(f"Saved: {filepath}")


# === 高级绘图函数 ===

def draw_flowchart(ax, steps, layout='horizontal', box_w=2.5, box_h=1.5,
                   gap=3.0, start_x=2.0, start_y=None, arrow_color=None):
    """
    绘制步骤流程图。

    steps: list of dict, 每个 dict 包含:
        - id: str, 步骤编号如 'S1'
        - title: str, 步骤标题
        - desc: str, 步骤描述（可选）
        - color: str, 背景色（可选，默认 light_blue）
        - edge: str, 边框色（可选，默认 secondary）

    layout: 'horizontal' | 'u_shape'
    返回 positions 列表。
    """
    default_color = COLORS['light_blue']
    default_edge = COLORS['secondary']
    if arrow_color is None:
        arrow_color = COLORS['secondary']

    if start_y is None:
        start_y = ax.get_ylim()[1] * 0.6

    positions = []

    if layout == 'horizontal':
        for i, step in enumerate(steps):
            x = start_x + i * gap
            y = start_y
            positions.append((x, y))
    elif layout == 'u_shape':
        n = len(steps)
        top_n = (n + 1) // 2
        bot_n = n - top_n
        top_y = start_y
        bot_y = start_y - box_h - gap * 0.8
        for i in range(top_n):
            positions.append((start_x + i * gap, top_y))
        for i in range(bot_n):
            positions.append((start_x + (top_n - 1 - i) * gap, bot_y))

    for i, step in enumerate(steps):
        x, y = positions[i]
        fc = step.get('color', default_color)
        ec = step.get('edge', default_edge)

        draw_box(ax, x, y, box_w, box_h, '', fc, edgecolor=ec, lw=1.8)
        draw_circle_label(ax, x - box_w/2 + 0.25, y + box_h/2 - 0.25,
                          step['id'], facecolor=ec)
        ax.text(x, y + 0.2, step['title'], ha='center', va='center',
                fontsize=11, color=COLORS['dark'], fontweight='bold')
        if step.get('desc'):
            ax.text(x, y - 0.3, step['desc'], ha='center', va='center',
                    fontsize=8.5, color='#555555')

    # Draw arrows between consecutive steps
    for i in range(len(steps) - 1):
        x1, y1 = positions[i]
        x2, y2 = positions[i + 1]
        if abs(y1 - y2) < 0.1:  # same row
            if x2 > x1:
                draw_arrow(ax, x1 + box_w/2 + 0.05, y1,
                           x2 - box_w/2 - 0.05, y2, color=arrow_color, lw=2)
            else:
                draw_arrow(ax, x1 - box_w/2 - 0.05, y1,
                           x2 + box_w/2 + 0.05, y2, color=arrow_color, lw=2)
        else:  # different row (U-shape turn)
            draw_arrow(ax, x1, y1 - box_h/2 - 0.05,
                       x2, y2 + box_h/2 + 0.05, color=arrow_color, lw=2)

    return positions


def draw_layered_arch(ax, layers, x=1.0, total_w=12.0):
    """
    绘制分层架构图。

    layers: list of dict (从上到下), 每个 dict 包含:
        - y: float, 层的 y 坐标
        - h: float, 层的高度
        - title: str, 层标题
        - bg: str, 背景色
        - edge: str, 边框色
        - items: list of dict, 每个 item 包含 text, color(可选), edge(可选)
    """
    for layer in layers:
        draw_layer_bg(ax, x, layer['y'], total_w, layer['h'],
                      layer['title'], layer['bg'], layer['edge'])
        if 'items' in layer:
            items = layer['items']
            n = len(items)
            item_w = min(2.5, (total_w - 1) / n - 0.3)
            item_gap = (total_w - 0.6) / n
            item_y = layer['y'] + layer['h'] / 2
            for i, item in enumerate(items):
                ix = x + 0.3 + item_gap * (i + 0.5)
                fc = item.get('color', COLORS['white'])
                ec = item.get('edge', layer['edge'])
                draw_box(ax, ix, item_y, item_w, 0.8, item['text'],
                         fc, edgecolor=ec, fontsize=9, lw=1.2)


def insert_images_to_cell(doc_path, table_idx, row_idx, col_idx, fig_map, width_cm=15):
    """
    将图片插入到 Word 文档指定表格单元格中。

    fig_map: dict, key=标记文字（如'附图1'）, value=图片文件路径
    在单元格中找到包含标记文字的段落，在其后插入图片。
    """
    from docx import Document
    from docx.shared import Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    doc = Document(doc_path)
    cell = doc.tables[table_idx].cell(row_idx, col_idx)
    tc = cell._tc

    for marker, img_path in fig_map.items():
        for p_elem in list(tc.iterchildren(qn('w:p'))):
            p_text = ''.join(node.text or '' for node in p_elem.iter(qn('w:t')))
            if marker in p_text:
                new_p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                jc = OxmlElement('w:jc')
                jc.set(qn('w:val'), 'center')
                pPr.append(jc)
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:before'), '120')
                spacing.set(qn('w:after'), '120')
                pPr.append(spacing)
                new_p.append(pPr)
                p_elem.addnext(new_p)

                new_para = Paragraph(new_p, cell)
                run = new_para.add_run()
                run.add_picture(img_path, width=Cm(width_cm))
                print(f"Inserted {marker} -> {img_path}")
                break

    doc.save(doc_path)
    print(f"Document saved: {doc_path}")
